from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

document=Document()

section = document.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2') 
img=input('Enter path of your Image please?  ')
document.add_picture(img, width=Inches(1.5))
first_name=input('What is your First name?  ')
last_name=input('What is your Last Name?  ')
document.add_heading(first_name+' '+last_name)

#about me
document.add_heading('About Me')
about_me=input('Tell me about yourself?  ')
document.add_paragraph(about_me)

#Education
document.add_heading('Education')
p=document.add_paragraph()
university = input('Enter University:  ')
department =input('Enter Department:  ')
from_date=input('From Date:  ')
to_date=input('To Date:  ')
p.add_run(university +'  ').bold=True
p.add_run(department +'  ').bold=True
p.add_run(from_date+' - '+to_date+'\n').italic =True
    
while True:
    has_more_experience=input('Do you have more education? (yes or no)  ')
    if has_more_experience.lower()=='yes':
        p=document.add_paragraph()
        university = input('Enter University:  ')
        department =input('Enter Department:  ')
        from_date=input('From Date:  ')
        to_date=input('To Date:  ')
        p.add_run(university +'  ').bold=True
        p.add_run(department +'  ').bold=True
        p.add_run(from_date+' - '+to_date+'\n').italic =True
    
    else:
        break

#Work Experience
document.add_heading('Work Experience')
p=document.add_paragraph()

company = input('Enter Company:  ')
from_date=input('From Date:  ')
to_date=input('To Date:  ')
p.add_run(company +'  ').bold=True
p.add_run(from_date+' - '+to_date+'\n').italic =True
experience_details=input('Describe your experience at'+' '+company+'  ')
p.add_run(experience_details)

while True:
    has_more_experience=input('Do you have more experience? (yes or no)  ')
    if has_more_experience.lower()=='yes':
        p=document.add_paragraph()
        company = input('Enter Company:  ')
        from_date=input('From Date:  ')
        to_date=input('To Date:  ')
        p.add_run(company +'  ').bold=True
        p.add_run( from_date + ' - ' + to_date+ '\n' ).italic=True
        experience_details=input('Describe your experience at'+' '+company+'  ')
        p.add_run(experience_details)
    
    else:
        break

#Skills 
document.add_heading('Skills')
s=document.add_paragraph()
s.style='List Bullet'
skills_details=input('Describe your Skills?  ')
s.add_run(experience_details)

while True:
    has_more_skills=input('Do you have more skills? (yes or no)  ')
    if has_more_skills.lower()=='yes':
        s=document.add_paragraph()
        s.style='List Bullet'
        skill=input('What is your skill?  ')
        s.add_run(skill)
    
    else:
        break

#Languages
document.add_heading('Languages')
s=document.add_paragraph()
s.style='List Bullet'
lan=input('What is your language?  ')
rate=input('Language Proficiency Ratio?  ')
s.add_run(lan + '  ' + rate)

while True:
    has_more_languages=input('Do you speak another language? (yes or no)  ')
    if has_more_languages.lower()=='yes':
        s=document.add_paragraph()
        s.style='List Bullet'
        lan=input('What is your language?  ')
        rate=input('Language Proficiency Ratio %?  ')
        s.add_run(lan + '  ' + rate)
    
    else:
        break
#contact
document.add_heading('Contact')
s=document.add_paragraph()
s.style='List Bullet'
phone_number=input('What is your phone?  ')
email=input('What is your email?  ')
s.add_run(phone_number+'\n')
s.add_run(email+'\n')

#footer
section=document.section[0]
footer=section.footer
p=footer.paragraph[0]
p.text="This cv generated using python. designedby hima hamod"
document.save('cv_'+first_name+'.docx')
